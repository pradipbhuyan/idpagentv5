import json
import re
import time
import base64
from io import BytesIO
from pathlib import Path
import tempfile
from datetime import datetime
import uuid
from difflib import SequenceMatcher

import pandas as pd
from docx import Document as DocxDocument
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage
from streamlit import session_state as st_state

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib import colors

MODEL_PRICING = {
    "gpt-4o-mini": {"input_per_1k": 0.00015, "output_per_1k": 0.0006},
    "gpt-4o": {"input_per_1k": 0.005, "output_per_1k": 0.015},
    "gpt-5": {"input_per_1k": 0.0, "output_per_1k": 0.0},
}

REQUIRED_RESUME_PLACEHOLDERS = [
    "{{name}}",
    "{{email}}",
    "{{phone}}",
    "{{location}}",
    "{{linkedin}}",
    "{{summary}}",
    "{{skills}}",
    "{{experience}}",
    "{{education}}",
    "{{certifications}}",
    "{{projects}}",
]

# ------------------------------
# METRICS
# ------------------------------
def ensure_metrics_state():
    if "metrics" not in st_state:
        st_state["metrics"] = {
            "tokens": 0,
            "input_tokens": 0,
            "output_tokens": 0,
            "cost": 0.0,
            "response_times": [],
            "calls": 0
        }

    if "doc_costs" not in st_state:
        st_state["doc_costs"] = {}


def get_current_metrics_snapshot():
    ensure_metrics_state()
    m = st_state["metrics"]
    return {
        "tokens": m.get("tokens", 0),
        "input_tokens": m.get("input_tokens", 0),
        "output_tokens": m.get("output_tokens", 0),
        "cost": m.get("cost", 0.0),
        "calls": m.get("calls", 0),
    }


def diff_metrics_snapshot(before, after):
    return {
        "tokens": after.get("tokens", 0) - before.get("tokens", 0),
        "input_tokens": after.get("input_tokens", 0) - before.get("input_tokens", 0),
        "output_tokens": after.get("output_tokens", 0) - before.get("output_tokens", 0),
        "cost": after.get("cost", 0.0) - before.get("cost", 0.0),
        "calls": after.get("calls", 0) - before.get("calls", 0),
    }


def get_model_pricing(model_name: str):
    return MODEL_PRICING.get(model_name, MODEL_PRICING.get("gpt-4o-mini"))


def invoke_llm_tracked(prompt: str):
    if "api_key" not in st_state:
        raise ValueError("Missing API key")

    model_name = st_state.get("model_choice", "gpt-4o-mini")
    
    llm = ChatOpenAI(
    model=model_name,
    temperature=0,
    api_key=st_state["api_key"],
    timeout=90,
    max_retries=1,
    )

    start = time.time()
    response = llm.invoke(prompt)
    duration = time.time() - start

    usage = getattr(response, "response_metadata", {}).get("token_usage", {}) or {}
    input_tokens = usage.get("prompt_tokens", 0)
    output_tokens = usage.get("completion_tokens", 0)

    if not input_tokens and not output_tokens:
        input_tokens = len(str(prompt)) // 4
        output_tokens = len(str(getattr(response, "content", ""))) // 4

    total_tokens = input_tokens + output_tokens
    pricing = get_model_pricing(model_name)
    input_cost = input_tokens * pricing["input_per_1k"] / 1000
    output_cost = output_tokens * pricing["output_per_1k"] / 1000
    total_cost = input_cost + output_cost

    ensure_metrics_state()
    m = st_state["metrics"]
    m["tokens"] += total_tokens
    m["input_tokens"] += input_tokens
    m["output_tokens"] += output_tokens
    m["cost"] += total_cost
    m["calls"] += 1
    m["response_times"].append(duration)

    doc = st_state.get("current_file") or "unknown"
    if doc not in st_state["doc_costs"]:
        st_state["doc_costs"][doc] = {"cost": 0.0, "tokens": 0}

    st_state["doc_costs"][doc]["cost"] += total_cost
    st_state["doc_costs"][doc]["tokens"] += total_tokens

    return response

# ------------------------------
# OCR / EXTRACTION QUALITY
# ------------------------------
def needs_ocr_fallback(text: str, min_chars: int = 120) -> bool:
    if not text:
        return True

    stripped = text.strip()
    if len(stripped) < min_chars:
        return True

    alnum_ratio = sum(ch.isalnum() for ch in stripped) / max(len(stripped), 1)
    if alnum_ratio < 0.2:
        return True

    lines = [ln.strip() for ln in stripped.splitlines() if ln.strip()]
    if len(lines) <= 2 and len(stripped) < 300:
        return True

    return False


def ocr_image_bytes_with_vlm(image_bytes: bytes, mime_type: str = "image/png") -> str:
    if "api_key" not in st_state:
        raise ValueError("Missing API key")

    model_name = st_state.get("model_choice", "gpt-4o-mini")

    llm = ChatOpenAI(
    model=model_name,
    temperature=0,
    api_key=st_state["api_key"],
    timeout=90,
    max_retries=1,
    )

    encoded = base64.b64encode(image_bytes).decode()
    msg = HumanMessage(
        content=[
            {
                "type": "text",
                "text": """Extract all visible text from this document image.

Rules:
- Output plain text only
- Preserve numbers, dates, amounts, and layout as much as possible
- Do not summarize
"""
            },
            {
                "type": "image_url",
                "image_url": {"url": f"data:{mime_type};base64,{encoded}"}
            }
        ]
    )

    start = time.time()
    response = llm.invoke([msg])
    duration = time.time() - start

    ensure_metrics_state()
    m = st_state["metrics"]
    content = getattr(response, "content", "") or ""
    input_tokens = 250
    output_tokens = len(str(content)) // 4
    pricing = get_model_pricing(model_name)
    total_cost = (input_tokens * pricing["input_per_1k"] / 1000) + (output_tokens * pricing["output_per_1k"] / 1000)

    m["tokens"] += input_tokens + output_tokens
    m["input_tokens"] += input_tokens
    m["output_tokens"] += output_tokens
    m["cost"] += total_cost
    m["calls"] += 1
    m["response_times"].append(duration)

    doc = st_state.get("current_file") or "unknown"
    if doc not in st_state["doc_costs"]:
        st_state["doc_costs"][doc] = {"cost": 0.0, "tokens": 0}
    st_state["doc_costs"][doc]["cost"] += total_cost
    st_state["doc_costs"][doc]["tokens"] += input_tokens + output_tokens

    return str(content).strip()


def extract_text_from_pdf_with_ocr_fallback(file_path: str):
    from langchain_community.document_loaders import PyPDFLoader

    docs = PyPDFLoader(file_path).load()
    raw_text = "\n".join([d.page_content for d in docs if getattr(d, "page_content", None)]).strip()

    if not needs_ocr_fallback(raw_text):
        return {
            "text": raw_text,
            "ocr_used": False,
            "extraction_mode": "native_pdf_text",
            "exception_reason": None,
        }

    try:
        import fitz
    except Exception:
        return {
            "text": raw_text,
            "ocr_used": False,
            "extraction_mode": "native_pdf_text_weak",
            "exception_reason": "OCR fallback needed but PyMuPDF is not available",
        }

    try:
        pdf = fitz.open(file_path)
        pages_text = []

        for page_index in range(len(pdf)):
            page = pdf.load_page(page_index)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_bytes = pix.tobytes("png")
            page_text = ocr_image_bytes_with_vlm(img_bytes, mime_type="image/png")
            if page_text:
                pages_text.append(f"Page {page_index + 1}\n{page_text}")

        ocr_text = "\n\n".join(pages_text).strip()

        if needs_ocr_fallback(ocr_text):
            return {
                "text": raw_text or ocr_text,
                "ocr_used": True,
                "extraction_mode": "pdf_ocr_attempted_weak",
                "exception_reason": "OCR fallback produced weak text",
            }

        return {
            "text": ocr_text,
            "ocr_used": True,
            "extraction_mode": "pdf_ocr_vlm",
            "exception_reason": None,
        }
    except Exception as e:
        return {
            "text": raw_text,
            "ocr_used": False,
            "extraction_mode": "native_pdf_text_weak",
            "exception_reason": f"OCR fallback failed: {str(e)}",
        }

# ------------------------------
# JSON / EXTRACTION
# ------------------------------
def safe_json_parse(text):
    if not text:
        return {}

    text = text.strip().replace("```json", "").replace("```", "").strip()

    try:
        return json.loads(text)
    except Exception:
        pass

    try:
        match = re.search(r"\{.*\}", text, re.DOTALL)
        if match:
            return json.loads(match.group())
    except Exception:
        pass

    try:
        text_fixed = re.sub(r",\s*}", "}", text)
        text_fixed = re.sub(r",\s*]", "]", text_fixed)
        return json.loads(text_fixed)
    except Exception:
        pass

    return {"raw_output": text}


def extract_structured_json(text, doc_type):
    clean_text = re.sub(r"[^\x00-\x7F]+", " ", text or "")
    clean_text = clean_text.replace("{", "").replace("}", "").strip()

    if "api_key" not in st_state:
        return {"error": "Missing API key"}

    if doc_type == "resume":
        prompt = f"""
You are a strict JSON generator.

Return ONLY valid JSON.
Do not add markdown.
Do not wrap in triple backticks.
Do not add explanation.

STRICT SCHEMA:
{{
  "name": "",
  "email": "",
  "phone": "",
  "location": "",
  "linkedin": "",
  "skills": [],
  "summary": "",
  "education": [
    {{
      "institution": "",
      "degree": "",
      "field_of_study": "",
      "start_date": "",
      "end_date": "",
      "graduation_date": "",
      "location": "",
      "details": []
    }}
  ],
  "experience": [
    {{
      "company": "",
      "role": "",
      "location": "",
      "start_date": "",
      "end_date": "",
      "is_current": false,
      "description": []
    }}
  ],
  "certifications": [
    {{
      "name": "",
      "issuer": "",
      "date": "",
      "expiry_date": ""
    }}
  ],
  "projects": [
    {{
      "name": "",
      "role": "",
      "start_date": "",
      "end_date": "",
      "description": []
    }}
  ]
}}

STRICT RULES:
- Extract ALL experience entries
- Extract ALL education entries
- Preserve ALL dates exactly
- Use empty strings for missing scalars
- Use empty arrays for missing lists

CV TEXT:
{clean_text[:12000]}
"""
    elif doc_type == "invoice":
        prompt = f"""
You are a strict JSON extractor.

Return ONLY valid JSON.
No markdown.
No explanation.

Extract invoice fields such as:
- vendor
- supplier
- invoice_number
- invoice_no
- invoice_date
- due_date
- currency
- subtotal
- tax
- total
- purchase_order
- line_items

DOCUMENT TEXT:
{clean_text[:12000]}
"""
    elif doc_type == "ticket":
        prompt = f"""
You are a strict JSON extractor.

Return ONLY valid JSON.
No markdown.
No explanation.

Extract travel ticket fields such as:
- traveler_name
- ticket_number
- booking_reference
- airline
- from
- to
- departure_date
- return_date
- amount
- currency
- class
- trip_type

DOCUMENT TEXT:
{clean_text[:12000]}
"""
    else:
        return {}

    try:
        response = invoke_llm_tracked(prompt).content.strip()
        response = response.replace("```json", "").replace("```", "").strip()
        parsed = safe_json_parse(response)

        if isinstance(parsed, list):
            merged = {}
            for item in parsed:
                if isinstance(item, dict):
                    merged.update(item)
            parsed = merged if merged else {"data": parsed}

        if not isinstance(parsed, dict):
            parsed = {"data": parsed}

        if doc_type == "resume":
            if not parsed.get("name"):
                try:
                    name_prompt = f"""
Extract only the candidate's full name from this resume text.
Return only the name.

{clean_text[:3000]}
"""
                    parsed["name"] = invoke_llm_tracked(name_prompt).content.strip()
                except Exception:
                    parsed["name"] = "Candidate"

            for field in ["name", "email", "phone", "location", "linkedin", "summary"]:
                parsed[field] = str(parsed.get(field, "") or "")

            parsed["skills"] = parsed.get("skills", []) if isinstance(parsed.get("skills", []), list) else []
            parsed["education"] = parsed.get("education", []) if isinstance(parsed.get("education", []), list) else []
            parsed["experience"] = parsed.get("experience", []) if isinstance(parsed.get("experience", []), list) else []
            parsed["certifications"] = parsed.get("certifications", []) if isinstance(parsed.get("certifications", []), list) else []
            parsed["projects"] = parsed.get("projects", []) if isinstance(parsed.get("projects", []), list) else []

        return parsed

    except Exception as e:
        return {"error": "LLM request failed", "details": str(e)[:300]}

# ------------------------------
# CONFIDENCE + VALIDATION
# ------------------------------
def confidence_label(score):
    if score >= 0.85:
        return "High"
    if score >= 0.6:
        return "Medium"
    return "Low"


def build_confidence_map(data, doc_type):
    if not isinstance(data, dict):
        return {}

    def score_scalar(value, strong=False):
        if value in [None, "", [], {}]:
            return {"score": 0.2, "label": "Low", "reason": "Missing or empty field"}
        if strong:
            score = 0.9
            reason = "Looks like an explicit field match"
        else:
            score = 0.7
            reason = "Extracted successfully but may need review"
        return {"score": score, "label": confidence_label(score), "reason": reason}

    confidence = {}

    if doc_type == "invoice":
        for field in ["vendor", "invoice_number", "invoice_date", "total", "currency", "due_date"]:
            val = data.get(field) or data.get(field.replace("invoice_number", "invoice_no"))
            confidence[field] = score_scalar(val, strong=field in ["invoice_number", "total"])

    elif doc_type == "ticket":
        for field in ["traveler_name", "ticket_number", "airline", "from", "to", "departure_date", "amount"]:
            confidence[field] = score_scalar(data.get(field), strong=field in ["ticket_number", "departure_date"])

    elif doc_type == "resume":
        for field in ["name", "email", "phone", "location", "summary"]:
            confidence[field] = score_scalar(data.get(field), strong=field in ["name", "email"])
        confidence["experience"] = score_scalar(data.get("experience"), strong=True)
        confidence["education"] = score_scalar(data.get("education"), strong=True)

    return confidence


def validate_document_data(data, doc_type):
    issues = []
    warnings = []

    if not isinstance(data, dict):
        return {"passed": False, "issues": ["No structured data available"], "warnings": []}

    if doc_type == "invoice":
        if not (data.get("vendor") or data.get("supplier")):
            issues.append("Vendor is missing")
        if not (data.get("invoice_number") or data.get("invoice_no")):
            issues.append("Invoice number is missing")
        if not data.get("invoice_date"):
            issues.append("Invoice date is missing")
        if not data.get("total"):
            issues.append("Total amount is missing")

    elif doc_type == "ticket":
        if not data.get("traveler_name"):
            issues.append("Traveler name is missing")
        if not data.get("ticket_number"):
            issues.append("Ticket number is missing")
        if not data.get("from") or not data.get("to"):
            issues.append("Route is incomplete")
        if not data.get("departure_date"):
            issues.append("Departure date is missing")
        if not data.get("amount"):
            warnings.append("Amount is missing")

    elif doc_type == "resume":
        if not data.get("name"):
            issues.append("Candidate name is missing")
        if not data.get("experience"):
            issues.append("Experience section is missing")
        if not data.get("education"):
            warnings.append("Education section is missing")
        if not data.get("skills"):
            warnings.append("Skills section is missing")

    return {"passed": len(issues) == 0, "issues": issues, "warnings": warnings}


def classify_exception(doc_type, text, validation, confidence, extraction_meta):
    if extraction_meta.get("exception_reason"):
        return extraction_meta["exception_reason"]

    if needs_ocr_fallback(text):
        return "No extractable text"

    if validation and not validation.get("passed", True):
        return "Validation failed"

    low_conf = [k for k, v in (confidence or {}).items() if v.get("label") == "Low"]
    if len(low_conf) >= 2:
        return "Low confidence"

    return None

# ------------------------------
# TEMPLATE MANAGER
# ------------------------------
def extract_docx_placeholders(template_file):
    if not template_file:
        return []

    try:
        if isinstance(template_file, bytes):
            doc = DocxDocument(BytesIO(template_file))
        elif hasattr(template_file, "read"):
            content = template_file.read()
            if hasattr(template_file, "seek"):
                template_file.seek(0)
            doc = DocxDocument(BytesIO(content))
        elif isinstance(template_file, str):
            doc = DocxDocument(template_file)
        else:
            return []
    except Exception:
        return []

    text_parts = []

    for para in doc.paragraphs:
        if para.text:
            text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        text_parts.append(para.text)

    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text:
                text_parts.append(para.text)
        for para in section.footer.paragraphs:
            if para.text:
                text_parts.append(para.text)

    all_text = "\n".join(text_parts)
    placeholders = sorted(set(re.findall(r"\{\{[^{}]+\}\}", all_text)))
    return placeholders


def validate_resume_template(template_file):
    found = extract_docx_placeholders(template_file)
    missing = [p for p in REQUIRED_RESUME_PLACEHOLDERS if p not in found]

    return {
        "valid": len(missing) == 0,
        "found_placeholders": found,
        "missing_placeholders": missing,
        "required_placeholders": REQUIRED_RESUME_PLACEHOLDERS,
    }

# ------------------------------
# RESUME
# ------------------------------
def generate_resume_summary(data):
    if "api_key" not in st_state:
        return "Summary not available"

    prompt = f"""
Create a professional resume summary in plain text.

STRICT RULES:
- No markdown
- Plain text only
- 4 to 6 concise lines
- Mention strengths, domain, and seniority
- Do not invent facts

CANDIDATE DATA:
{json.dumps(data, ensure_ascii=False)}
"""
    try:
        return invoke_llm_tracked(prompt).content.strip()
    except Exception:
        return "Summary not available"


def build_resume(data, template_file):
    def safe_str(value):
        return "" if value is None else str(value)

    def format_date_range(start_date, end_date):
        start_date = safe_str(start_date).strip()
        end_date = safe_str(end_date).strip()
        if start_date and end_date:
            return f"{start_date} - {end_date}"
        if start_date:
            return start_date
        if end_date:
            return end_date
        return ""

    def format_skills(skills):
        if not isinstance(skills, list) or not skills:
            return ""
        return ", ".join(str(s).strip() for s in skills if str(s).strip())

    def format_experience(experience):
        if not isinstance(experience, list) or not experience:
            return ""
        lines = []
        for exp in experience:
            if not isinstance(exp, dict):
                continue
            title = " - ".join([safe_str(exp.get("role")).strip(), safe_str(exp.get("company")).strip()]).strip(" -")
            date_text = format_date_range(exp.get("start_date"), exp.get("end_date"))
            location = safe_str(exp.get("location")).strip()
            first = " ".join([p for p in [title, f"({date_text})" if date_text else "", location] if p]).strip()
            if first:
                lines.append(first)
            for item in exp.get("description", []):
                item = safe_str(item).strip()
                if item:
                    lines.append(f"- {item}")
            lines.append("")
        return "\n".join(lines).strip()

    def format_education(education):
        if not isinstance(education, list) or not education:
            return ""
        lines = []
        for edu in education:
            if not isinstance(edu, dict):
                continue
            first = " - ".join(
                [p for p in [safe_str(edu.get("degree")).strip(), safe_str(edu.get("institution")).strip()] if p]
            )
            if first:
                lines.append(first)
            date_text = safe_str(edu.get("graduation_date")).strip() or format_date_range(
                edu.get("start_date"), edu.get("end_date")
            )
            if date_text or edu.get("location"):
                lines.append(", ".join([p for p in [date_text, safe_str(edu.get("location")).strip()] if p]))
            for item in edu.get("details", []):
                item = safe_str(item).strip()
                if item:
                    lines.append(f"- {item}")
            lines.append("")
        return "\n".join(lines).strip()

    def format_certifications(certifications):
        if not isinstance(certifications, list) or not certifications:
            return ""
        lines = []
        for cert in certifications:
            if not isinstance(cert, dict):
                continue
            name = safe_str(cert.get("name")).strip()
            issuer = safe_str(cert.get("issuer")).strip()
            date = safe_str(cert.get("date")).strip()
            text = " - ".join([p for p in [name, issuer] if p])
            if date:
                text = f"{text} ({date})" if text else date
            if text:
                lines.append(text)
        return "\n".join(lines).strip()

    def format_projects(projects):
        if not isinstance(projects, list) or not projects:
            return ""
        lines = []
        for proj in projects:
            if not isinstance(proj, dict):
                continue
            title = " - ".join([safe_str(proj.get("name")).strip(), safe_str(proj.get("role")).strip()]).strip(" -")
            date_text = format_date_range(proj.get("start_date"), proj.get("end_date"))
            first = " ".join([p for p in [title, f"({date_text})" if date_text else ""] if p]).strip()
            if first:
                lines.append(first)
            for item in proj.get("description", []):
                item = safe_str(item).strip()
                if item:
                    lines.append(f"- {item}")
            lines.append("")
        return "\n".join(lines).strip()

    def replace_placeholders_in_paragraph(paragraph, placeholders):
        for key, value in placeholders.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    def replace_placeholders(doc, placeholders):
        for para in doc.paragraphs:
            replace_placeholders_in_paragraph(para, placeholders)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_placeholders_in_paragraph(para, placeholders)
        for section in doc.sections:
            for para in section.header.paragraphs:
                replace_placeholders_in_paragraph(para, placeholders)
            for para in section.footer.paragraphs:
                replace_placeholders_in_paragraph(para, placeholders)

    summary = data.get("summary") or generate_resume_summary(data)

    if not template_file:
        raise ValueError("No template file provided")

    try:
        if isinstance(template_file, bytes):
            doc = DocxDocument(BytesIO(template_file))
        elif hasattr(template_file, "read"):
            content = template_file.read()
            if hasattr(template_file, "seek"):
                template_file.seek(0)
            doc = DocxDocument(BytesIO(content))
        elif isinstance(template_file, str):
            doc = DocxDocument(template_file)
        else:
            raise TypeError(f"Unsupported template_file type: {type(template_file)}")
    except Exception as e:
        raise RuntimeError(f"Template load failed: {e}")

    placeholders = {
        "{{name}}": safe_str(data.get("name", "")),
        "{{email}}": safe_str(data.get("email", "")),
        "{{phone}}": safe_str(data.get("phone", "")),
        "{{location}}": safe_str(data.get("location", "")),
        "{{linkedin}}": safe_str(data.get("linkedin", "")),
        "{{summary}}": safe_str(summary),
        "{{skills}}": format_skills(data.get("skills", [])),
        "{{experience}}": format_experience(data.get("experience", [])),
        "{{education}}": format_education(data.get("education", [])),
        "{{certifications}}": format_certifications(data.get("certifications", [])),
        "{{projects}}": format_projects(data.get("projects", [])),
    }

    replace_placeholders(doc, placeholders)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ------------------------------
# DUPLICATE DETECTION
# ------------------------------
def normalize_text_for_match(value):
    if value is None:
        return ""
    text = str(value).strip().lower()
    text = re.sub(r"\s+", " ", text)
    return text


def similarity_score(a, b):
    a_norm = normalize_text_for_match(a)
    b_norm = normalize_text_for_match(b)
    if not a_norm or not b_norm:
        return 0.0
    return SequenceMatcher(None, a_norm, b_norm).ratio()


def generate_duplicate_key(doc_type, data):
    if not isinstance(data, dict):
        return None

    if doc_type == "invoice":
        vendor = data.get("vendor") or data.get("supplier") or ""
        invoice_no = data.get("invoice_number") or data.get("invoice_no") or ""
        total = data.get("total") or ""
        invoice_date = data.get("invoice_date") or ""
        return f"invoice|{normalize_text_for_match(vendor)}|{normalize_text_for_match(invoice_no)}|{normalize_text_for_match(total)}|{normalize_text_for_match(invoice_date)}"

    if doc_type == "ticket":
        traveler = data.get("traveler_name") or ""
        ticket_no = data.get("ticket_number") or ""
        route = f"{data.get('from', '')}-{data.get('to', '')}"
        departure_date = data.get("departure_date") or ""
        return f"ticket|{normalize_text_for_match(traveler)}|{normalize_text_for_match(ticket_no)}|{normalize_text_for_match(route)}|{normalize_text_for_match(departure_date)}"

    if doc_type == "resume":
        name = data.get("name") or ""
        email = data.get("email") or ""
        phone = data.get("phone") or ""
        return f"resume|{normalize_text_for_match(name)}|{normalize_text_for_match(email)}|{normalize_text_for_match(phone)}"

    return None


def detect_duplicate_document(new_doc_type, new_data, existing_results):
    new_key = generate_duplicate_key(new_doc_type, new_data)
    if not new_key or not existing_results:
        return {
            "is_duplicate": False,
            "match_file": None,
            "reason": None,
            "score": 0.0,
        }

    for item in existing_results:
        existing_doc_type = item.get("doc_type")
        existing_data = item.get("review_data") or {}
        if existing_doc_type != new_doc_type:
            continue

        existing_key = generate_duplicate_key(existing_doc_type, existing_data)
        if not existing_key:
            continue

        if new_key == existing_key:
            return {
                "is_duplicate": True,
                "match_file": item.get("file_name"),
                "reason": "Exact duplicate key match",
                "score": 1.0,
            }

        score = similarity_score(new_key, existing_key)
        if score >= 0.92:
            return {
                "is_duplicate": True,
                "match_file": item.get("file_name"),
                "reason": "Near-duplicate structured match",
                "score": round(score, 3),
            }

    return {
        "is_duplicate": False,
        "match_file": None,
        "reason": None,
        "score": 0.0,
    }

# ------------------------------
# JD RANKING
# ------------------------------
def score_resume_against_jd(resume_data, jd_text):
    if not isinstance(resume_data, dict) or not jd_text:
        return {
            "candidate_name": "Unknown",
            "overall_score": 0,
            "skills_score": 0,
            "experience_score": 0,
            "education_score": 0,
            "matched_skills": [],
            "missing_skills": [],
            "strengths": [],
            "gaps": ["Insufficient input"],
            "recommendation": "Weak Fit"
        }

    prompt = f"""
You are a strict resume-job description matching assistant.

Compare the resume against the job description and return ONLY valid JSON.

Return JSON with this schema:
{{
  "candidate_name": "",
  "overall_score": 0,
  "skills_score": 0,
  "experience_score": 0,
  "education_score": 0,
  "matched_skills": [],
  "missing_skills": [],
  "strengths": [],
  "gaps": [],
  "recommendation": ""
}}

Rules:
- All scores must be integers from 0 to 100
- recommendation must be one of:
  "Strong Fit", "Moderate Fit", "Weak Fit"
- matched_skills and missing_skills should be concise
- strengths and gaps should be concise bullets
- use only the given data
- do not invent missing experience
- candidate_name should come from the resume if available

JOB DESCRIPTION:
{jd_text[:8000]}

RESUME DATA:
{json.dumps(resume_data, ensure_ascii=False)[:12000]}
"""

    try:
        response = invoke_llm_tracked(prompt).content.strip()
        parsed = safe_json_parse(response)

        if not isinstance(parsed, dict):
            parsed = {}

        parsed["candidate_name"] = str(parsed.get("candidate_name") or resume_data.get("name") or "Unknown")
        parsed["overall_score"] = int(parsed.get("overall_score", 0) or 0)
        parsed["skills_score"] = int(parsed.get("skills_score", 0) or 0)
        parsed["experience_score"] = int(parsed.get("experience_score", 0) or 0)
        parsed["education_score"] = int(parsed.get("education_score", 0) or 0)
        parsed["matched_skills"] = parsed.get("matched_skills", []) if isinstance(parsed.get("matched_skills", []), list) else []
        parsed["missing_skills"] = parsed.get("missing_skills", []) if isinstance(parsed.get("missing_skills", []), list) else []
        parsed["strengths"] = parsed.get("strengths", []) if isinstance(parsed.get("strengths", []), list) else []
        parsed["gaps"] = parsed.get("gaps", []) if isinstance(parsed.get("gaps", []), list) else []
        parsed["recommendation"] = str(parsed.get("recommendation") or "Moderate Fit")

        parsed["overall_score"] = max(0, min(100, parsed["overall_score"]))
        parsed["skills_score"] = max(0, min(100, parsed["skills_score"]))
        parsed["experience_score"] = max(0, min(100, parsed["experience_score"]))
        parsed["education_score"] = max(0, min(100, parsed["education_score"]))

        return parsed

    except Exception:
        return {
            "candidate_name": str(resume_data.get("name") or "Unknown"),
            "overall_score": 0,
            "skills_score": 0,
            "experience_score": 0,
            "education_score": 0,
            "matched_skills": [],
            "missing_skills": [],
            "strengths": [],
            "gaps": ["Scoring failed"],
            "recommendation": "Weak Fit"
        }

# ------------------------------
# CONCUR MOCK
# ------------------------------
def send_to_concur(doc_type, data, mode="mock"):
    payload = {"type": doc_type, "data": data}

    if doc_type == "invoice":
        try:
            payload["line_items"] = json_to_kv_dataframe(data).to_dict(orient="records")
        except Exception:
            payload["line_items"] = []

    now_utc = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
    short_id = uuid.uuid4().hex[:8].upper()
    batch_id = f"CCB-{datetime.utcnow().strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"
    endpoint = "Expense Entry Import API" if doc_type == "invoice" else "Travel Request / Expense Entry API"

    if mode == "mock":
        return {
            "status": "submitted",
            "mode": "mock",
            "message": f"{doc_type.title()} submitted to Concur mock gateway",
            "submission_id": f"SUB-{short_id}",
            "batch_id": batch_id,
            "document_id": f"{doc_type[:3].upper()}-{uuid.uuid4().hex[:10].upper()}",
            "submitted_at": now_utc,
            "endpoint": endpoint,
            "processing_state": "Queued for downstream validation",
            "next_status": "Expected to transition to Accepted or Rejected after validation",
            "payload": payload
        }

    return {
        "status": "submitted",
        "mode": "real",
        "message": f"{doc_type.title()} submitted to Concur",
        "submission_id": f"SUB-{short_id}",
        "batch_id": batch_id,
        "document_id": f"{doc_type[:3].upper()}-{uuid.uuid4().hex[:10].upper()}",
        "submitted_at": now_utc,
        "endpoint": endpoint,
        "processing_state": "Accepted by Concur endpoint",
        "next_status": "Awaiting downstream processing",
        "payload": payload
    }

# ------------------------------
# MISC
# ------------------------------
def save_temp_file(uploaded_file):
    suffix = Path(uploaded_file.name).suffix
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getvalue())
        return tmp.name


def detect_document_type(text):
    if "api_key" not in st_state:
        return "other"

    prompt = f"""
Classify document into ONE label:

resume
invoice
receipt
report
ticket
other

Return only the label.

{text[:2000]}
"""
    try:
        raw = invoke_llm_tracked(prompt).content.lower().strip()
    except Exception:
        return "other"

    labels = ["resume", "invoice", "receipt", "report", "ticket", "other"]
    for label in labels:
        if label == raw:
            return label
    for label in labels:
        if label in raw:
            return label
    return "other"


def json_to_kv_dataframe(data):
    rows = []

    def flatten(prefix, obj):
        if isinstance(obj, dict):
            for k, v in obj.items():
                flatten(f"{prefix}.{k}" if prefix else k, v)
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                flatten(f"{prefix}[{i}]", item)
        else:
            rows.append({
                "Field": prefix,
                "Value": json.dumps(obj) if isinstance(obj, (dict, list)) else str(obj)
            })

    flatten("", data if data is not None else {})
    return pd.DataFrame(rows)


def generate_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="data")
    return output.getvalue()

def _safe_list(value):
    return value if isinstance(value, list) else []


def _safe_dict(value):
    return value if isinstance(value, dict) else {}


def generate_recruiter_questions_from_jd(jd_text: str):
    jd_text = (jd_text or "").strip()

    if not jd_text:
        return [
            {
                "question": "Can you walk us through your most recent relevant project?",
                "expected_answer": "Candidate should clearly explain ownership, stack, outcomes, and delivery impact."
            },
            {
                "question": "Are you comfortable with hybrid work expectations and role-specific shift needs?",
                "expected_answer": "Candidate should confirm location, shift flexibility, and return-to-office readiness."
            }
        ]

    prompt = f"""
You are creating recruiter screening questions from a job description.

Return ONLY valid JSON in this exact format:
{{
  "questions": [
    {{
      "question": "",
      "expected_answer": ""
    }}
  ]
}}

Rules:
- Return 6 to 8 questions
- Questions should be recruiter-friendly, not too technical
- Focus on role fit, delivery ownership, location/work model, domain relevance, notice period, stakeholder handling, and authenticity checks
- expected_answer should be concise and practical
- No markdown
- No explanation

JOB DESCRIPTION:
{jd_text[:7000]}
"""
    try:
        response = invoke_llm_tracked(prompt).content.strip()
        parsed = safe_json_parse(response)
        questions = parsed.get("questions", []) if isinstance(parsed, dict) else []
        clean_questions = []

        for q in questions:
            if isinstance(q, dict):
                question = str(q.get("question", "")).strip()
                expected = str(q.get("expected_answer", "")).strip()
                if question:
                    clean_questions.append({
                        "question": question,
                        "expected_answer": expected or "-"
                    })

        if clean_questions:
            return clean_questions
    except Exception:
        pass

    return [
        {
            "question": "Can you explain your most recent role and your exact delivery responsibilities?",
            "expected_answer": "Candidate should clearly describe scope, role ownership, technology environment, and measurable outcomes."
        },
        {
            "question": "How closely does your experience align with this JD’s mandatory stack and domain expectations?",
            "expected_answer": "Candidate should map prior work to the required stack and honestly clarify any missing domain exposure."
        },
        {
            "question": "Are you currently available for the required location/work model and notice period expectations?",
            "expected_answer": "Candidate should confirm current location, hybrid readiness, shift flexibility, and notice details."
        }
    ]


def summarize_candidate_for_assessment(resume_data, ranking_data, file_name=""):
    resume_data = _safe_dict(resume_data)
    ranking_data = _safe_dict(ranking_data)

    name = (
        ranking_data.get("candidate_name")
        or resume_data.get("name")
        or file_name
        or "Candidate"
    )

    experience = _safe_list(resume_data.get("experience"))
    education = _safe_list(resume_data.get("education"))
    certifications = _safe_list(resume_data.get("certifications"))
    skills = _safe_list(resume_data.get("skills"))

    total_experience_entries = len(experience)
    total_education_entries = len(education)
    total_certifications = len(certifications)

    current_role = ""
    if experience:
        first_exp = experience[0] if isinstance(experience[0], dict) else {}
        current_role = " - ".join(
            [str(first_exp.get("role", "") or "").strip(), str(first_exp.get("company", "") or "").strip()]
        ).strip(" -")

    overall_score = int(ranking_data.get("overall_score", 0) or 0)
    recommendation = str(ranking_data.get("recommendation", "Moderate Fit"))

    if recommendation == "Strong Fit":
        shortlist_label = "Primary"
    elif recommendation == "Moderate Fit":
        shortlist_label = "Backup"
    else:
        shortlist_label = "Hold"

    return {
        "candidate_name": name,
        "file_name": file_name,
        "overall_score": overall_score,
        "skills_score": int(ranking_data.get("skills_score", 0) or 0),
        "experience_score": int(ranking_data.get("experience_score", 0) or 0),
        "education_score": int(ranking_data.get("education_score", 0) or 0),
        "recommendation": recommendation,
        "shortlist_label": shortlist_label,
        "matched_skills": _safe_list(ranking_data.get("matched_skills")),
        "missing_skills": _safe_list(ranking_data.get("missing_skills")),
        "strengths": _safe_list(ranking_data.get("strengths")),
        "gaps": _safe_list(ranking_data.get("gaps")),
        "location": resume_data.get("location", ""),
        "email": resume_data.get("email", ""),
        "phone": resume_data.get("phone", ""),
        "current_role": current_role,
        "skills": skills,
        "experience_count": total_experience_entries,
        "education_count": total_education_entries,
        "certification_count": total_certifications,
        "fitment_progress": f"{overall_score}%",
    }


def generate_consolidated_assessment_data(batch_results, jd_text, jd_rankings):
    resume_items = [
        item for item in (batch_results or [])
        if item.get("doc_type") == "resume" and item.get("review_data")
    ]

    ranking_map = {}
    for row in jd_rankings or []:
        key = row.get("file_name") or row.get("candidate_name")
        if key:
            ranking_map[key] = row

    candidates = []
    for item in resume_items:
        file_name = item.get("file_name", "")
        resume_data = item.get("review_data") or {}
        ranking_data = ranking_map.get(file_name, {})

        if not ranking_data:
            candidate_name = resume_data.get("name")
            if candidate_name:
                ranking_data = ranking_map.get(candidate_name, {})

        candidate_summary = summarize_candidate_for_assessment(
            resume_data=resume_data,
            ranking_data=ranking_data,
            file_name=file_name
        )
        candidates.append(candidate_summary)

    candidates = sorted(candidates, key=lambda x: x.get("overall_score", 0), reverse=True)

    highest_score = candidates[0]["overall_score"] if candidates else 0
    lowest_score = candidates[-1]["overall_score"] if candidates else 0

    primary = [c["candidate_name"] for c in candidates if c.get("shortlist_label") == "Primary"]
    backup = [c["candidate_name"] for c in candidates if c.get("shortlist_label") == "Backup"]
    hold = [c["candidate_name"] for c in candidates if c.get("shortlist_label") == "Hold"]

    if candidates:
        top_name = candidates[0]["candidate_name"]
        takeaway = (
            f"The candidate pool includes {len(candidates)} resume(s). "
            f"{top_name} currently leads based on evidence-backed JD fit, with strongest visible alignment across the extracted profiles."
        )
    else:
        takeaway = "No resume candidates were available for assessment."

    recruiter_questions = generate_recruiter_questions_from_jd(jd_text)

    return {
        "title": "DetailedAssesment",
        "executive_summary": {
            "analysis_date": datetime.now().strftime("%B %d, %Y"),
            "total_candidates": len(candidates),
            "top_match_range": f"{highest_score} → {lowest_score}" if candidates else "N/A",
            "recommended_action": (
                f"{len(primary)} primary / {len(backup)} backup"
                if candidates else "No recommendation"
            ),
            "jd_summary": jd_text[:1800] if jd_text else "No JD provided",
            "executive_takeaway": takeaway,
        },
        "candidates": candidates,
        "final_summary": {
            "primary_candidates": primary,
            "backup_candidates": backup,
            "hold_candidates": hold,
        },
        "recruiter_questions": recruiter_questions,
    }


def _score_badge_color(score: int):
    if score >= 85:
        return colors.HexColor("#138a36")
    if score >= 70:
        return colors.HexColor("#d98e04")
    return colors.HexColor("#b42318")


def build_consolidated_assessment_pdf(report_data):
    report_data = _safe_dict(report_data)
    executive = _safe_dict(report_data.get("executive_summary"))
    candidates = _safe_list(report_data.get("candidates"))
    final_summary = _safe_dict(report_data.get("final_summary"))
    recruiter_questions = _safe_list(report_data.get("recruiter_questions"))

    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Heading1"],
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#1f1f1f"),
        spaceAfter=10,
        alignment=TA_LEFT
    )
    section_style = ParagraphStyle(
        "SectionTitle",
        parent=styles["Heading2"],
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#1f1f1f"),
        spaceBefore=10,
        spaceAfter=8,
        alignment=TA_LEFT
    )
    body_style = ParagraphStyle(
        "BodySmall",
        parent=styles["BodyText"],
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#333333"),
        alignment=TA_LEFT
    )
    small_center = ParagraphStyle(
        "SmallCenter",
        parent=body_style,
        alignment=TA_CENTER
    )

    story = []

    story.append(Paragraph("Enhanced Candidate Screening Report", title_style))
    story.append(Paragraph(f"Analysis Date: {executive.get('analysis_date', '-')}", body_style))
    story.append(Paragraph("Confidential • Internal Recruiter Use Only", body_style))
    story.append(Spacer(1, 0.15 * inch))

    # Score cards
    card_data = [[
        Paragraph(f"<b>Candidates Evaluated</b><br/>{executive.get('total_candidates', 0)}", small_center),
        Paragraph(f"<b>Top-Match Range</b><br/>{executive.get('top_match_range', '-')}", small_center),
        Paragraph(f"<b>Recommended Action</b><br/>{executive.get('recommended_action', '-')}", small_center),
    ]]
    card_table = Table(card_data, colWidths=[1.8 * inch, 1.8 * inch, 2.0 * inch])
    card_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#eef6ff")),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#c8d7eb")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#c8d7eb")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(card_table)
    story.append(Spacer(1, 0.15 * inch))

    story.append(Paragraph("1. Executive Summary", section_style))
    story.append(Paragraph("<b>JD Summary</b>", body_style))
    story.append(Paragraph(str(executive.get("jd_summary", "-")).replace("\n", "<br/>"), body_style))
    story.append(Spacer(1, 0.08 * inch))
    story.append(Paragraph("<b>Executive Takeaway</b>", body_style))
    story.append(Paragraph(str(executive.get("executive_takeaway", "-")), body_style))
    story.append(Spacer(1, 0.15 * inch))

    story.append(Paragraph("2. Candidate Ranking & Shortlist Decision", section_style))
    rank_rows = [[
        Paragraph("<b>Rank</b>", body_style),
        Paragraph("<b>Candidate</b>", body_style),
        Paragraph("<b>Overall</b>", body_style),
        Paragraph("<b>Recommendation</b>", body_style),
        Paragraph("<b>Fitment Progress</b>", body_style),
    ]]
    for idx, c in enumerate(candidates, start=1):
        rank_rows.append([
            Paragraph(f"#{idx}", body_style),
            Paragraph(c.get("candidate_name", "-"), body_style),
            Paragraph(str(c.get("overall_score", 0)), body_style),
            Paragraph(c.get("shortlist_label", "-"), body_style),
            Paragraph(c.get("fitment_progress", "-"), body_style),
        ])
    rank_table = Table(rank_rows, colWidths=[0.55 * inch, 2.2 * inch, 0.85 * inch, 1.25 * inch, 1.15 * inch])
    rank_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f2f4f7")),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#d0d5dd")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e4e7ec")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(rank_table)
    story.append(Spacer(1, 0.15 * inch))

    for idx, c in enumerate(candidates, start=1):
        story.append(PageBreak())
        story.append(Paragraph(f"{idx}. {c.get('candidate_name', '-')}", section_style))

        score_color = _score_badge_color(int(c.get("overall_score", 0)))
        score_box = Table([[
            Paragraph(
                f"<font color='white'><b>Overall Score</b><br/>{c.get('overall_score', 0)}</font>",
                small_center
            )
        ]], colWidths=[1.2 * inch], rowHeights=[0.65 * inch])
        score_box.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), score_color),
            ("BOX", (0, 0), (-1, -1), 0.5, score_color),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))

        info_box = Table([[
            Paragraph(
                f"<b>File:</b> {c.get('file_name', '-')}<br/>"
                f"<b>Current Role:</b> {c.get('current_role', '-') or '-'}<br/>"
                f"<b>Location:</b> {c.get('location', '-') or '-'}<br/>"
                f"<b>Recommendation:</b> {c.get('recommendation', '-')}",
                body_style
            )
        ]], colWidths=[4.8 * inch])
        info_box.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#d0d5dd")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))

        header_table = Table([[info_box, score_box]], colWidths=[4.9 * inch, 1.3 * inch])
        story.append(header_table)
        story.append(Spacer(1, 0.1 * inch))

        metric_rows = [[
            Paragraph("<b>Skills</b>", small_center),
            Paragraph("<b>Experience</b>", small_center),
            Paragraph("<b>Education</b>", small_center),
            Paragraph("<b>Matched Skills</b>", small_center),
            Paragraph("<b>Missing Skills</b>", small_center),
        ], [
            Paragraph(str(c.get("skills_score", 0)), small_center),
            Paragraph(str(c.get("experience_score", 0)), small_center),
            Paragraph(str(c.get("education_score", 0)), small_center),
            Paragraph(", ".join(c.get("matched_skills", [])) or "-", body_style),
            Paragraph(", ".join(c.get("missing_skills", [])) or "-", body_style),
        ]]
        metric_table = Table(metric_rows, colWidths=[0.85 * inch, 0.95 * inch, 0.9 * inch, 2.0 * inch, 1.6 * inch])
        metric_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f2f4f7")),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#d0d5dd")),
            ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e4e7ec")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(metric_table)
        story.append(Spacer(1, 0.1 * inch))

        story.append(Paragraph("<b>Strengths</b>", body_style))
        strengths = c.get("strengths", [])
        if strengths:
            for s in strengths:
                story.append(Paragraph(f"• {s}", body_style))
        else:
            story.append(Paragraph("-", body_style))

        story.append(Spacer(1, 0.04 * inch))
        story.append(Paragraph("<b>Gaps / Risks</b>", body_style))
        gaps = c.get("gaps", [])
        if gaps:
            for g in gaps:
                story.append(Paragraph(f"• {g}", body_style))
        else:
            story.append(Paragraph("-", body_style))

        story.append(Spacer(1, 0.06 * inch))
        story.append(Paragraph(
            f"<b>Profile Indicators</b><br/>"
            f"Experience Entries: {c.get('experience_count', 0)}<br/>"
            f"Education Entries: {c.get('education_count', 0)}<br/>"
            f"Certifications: {c.get('certification_count', 0)}",
            body_style
        ))

    story.append(PageBreak())
    story.append(Paragraph("Final Consolidated Summary", section_style))

    shortlist_data = [
        [Paragraph("<b>Primary</b>", body_style), Paragraph(", ".join(final_summary.get("primary_candidates", [])) or "-", body_style)],
        [Paragraph("<b>Backup</b>", body_style), Paragraph(", ".join(final_summary.get("backup_candidates", [])) or "-", body_style)],
        [Paragraph("<b>Hold</b>", body_style), Paragraph(", ".join(final_summary.get("hold_candidates", [])) or "-", body_style)],
    ]
    shortlist_table = Table(shortlist_data, colWidths=[1.2 * inch, 5.2 * inch])
    shortlist_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f2f4f7")),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#d0d5dd")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e4e7ec")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(shortlist_table)
    story.append(Spacer(1, 0.15 * inch))

    story.append(Paragraph("Recruiter Screening Questions", section_style))
    q_rows = [[
        Paragraph("<b>Question</b>", body_style),
        Paragraph("<b>Expected Answer</b>", body_style),
    ]]
    for q in recruiter_questions:
        q_rows.append([
            Paragraph(str(q.get("question", "-")), body_style),
            Paragraph(str(q.get("expected_answer", "-")), body_style),
        ])
    q_table = Table(q_rows, colWidths=[3.0 * inch, 3.3 * inch])
    q_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f2f4f7")),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#d0d5dd")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e4e7ec")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(q_table)

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()
